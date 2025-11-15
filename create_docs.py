"""Generate comprehensive documentation for NSE Stock Performance Tracker"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('NSE Stock Performance Tracker', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('Complete Code Documentation & Architecture Guide')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc = ['1. Project Overview', '2. Architecture', '3. File Structure', '4. Module Explanations',
       '5. Data Flow', '6. Key Features', '7. Performance Optimizations', '8. Usage Guide']
for item in toc:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 1. Project Overview
doc.add_heading('1. Project Overview', 1)
doc.add_paragraph('The NSE Stock Performance Tracker is a Streamlit web application for tracking Indian stock market performance with real-time data, caching, and advanced visualizations.')

doc.add_heading('Key Capabilities:', 2)
caps = ['Track Nifty 50 and Nifty Next 50', 'Real-time market indices', 
        '1-week to 3-month performance metrics', 'Live rolling ticker', 'Custom stock lists',
        'Persistent caching (6-hour expiry)', 'Parallel data fetching', 'Commodity prices']
for cap in caps:
    doc.add_paragraph(cap, style='List Bullet')

doc.add_page_break()

# 2. Architecture
doc.add_heading('2. Architecture & Design', 1)
doc.add_paragraph('Modular architecture with clear separation of concerns:')
modules = [
    ('config.py', 'Constants, fallback data, CSS styling'),
    ('app.py', 'Main orchestration and business logic'),
    ('data_fetchers.py', 'External API calls and data retrieval'),
    ('ui_components.py', 'Streamlit UI rendering'),
    ('utils.py', 'Helper functions for formatting'),
    ('cache_manager.py', 'Persistent caching with Pickle'),
    ('file_manager.py', 'File I/O for custom lists')
]
for mod, desc in modules:
    p = doc.add_paragraph()
    p.add_run(mod + ': ').bold = True
    p.add_run(desc)

doc.add_page_break()

# 3. File Structure
doc.add_heading('3. File Structure', 1)
structure = """windsurf-project/
├── app.py                    # Main entry point
├── config.py                 # Configuration
├── data_fetchers.py          # Data fetching
├── ui_components.py          # UI rendering
├── utils.py                  # Utilities
├── cache_manager.py          # Cache management
├── file_manager.py           # File operations
├── requirements.txt          # Dependencies
├── cache/                    # Cache storage
└── saved_stock_lists/        # Custom lists"""
doc.add_paragraph(structure, style='Normal')

doc.add_page_break()

# 4. Module Explanations
doc.add_heading('4. Module-by-Module Explanation', 1)

# config.py
doc.add_heading('4.1 config.py', 2)
doc.add_paragraph('Centralizes configuration, constants, and styling.')
doc.add_heading('Key Components:', 3)
config_items = [
    ('SAVED_LISTS_DIR', 'Directory for user stock lists'),
    ('ITEMS_PER_PAGE', 'Pagination size (10)'),
    ('FALLBACK_NIFTY_50', 'Hardcoded Nifty 50 stocks'),
    ('FALLBACK_NIFTY_NEXT_50', 'Hardcoded Nifty Next 50'),
    ('INDICES_ROW1', 'Major indices (Nifty, Sensex, etc.)'),
    ('INDICES_ROW2', 'Sectoral indices (IT, Pharma, etc.)'),
    ('COMMODITIES', 'Commodity tickers (Oil, Gold, BTC)'),
    ('CUSTOM_CSS', 'Dark theme styling (440+ lines)'),
    ('METRIC_CSS', 'Metric component styling')
]
for item, desc in config_items:
    p = doc.add_paragraph()
    p.add_run(item + ': ').bold = True
    p.add_run(desc)

doc.add_heading('CSS Features:', 3)
css_features = ['Dark theme (#1e1e1e)', 'Responsive fonts (13-14px)', 'Custom sidebar',
                'Rolling ticker animation', 'Color-coded percentages', 'Hover effects']
for f in css_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# app.py
doc.add_heading('4.2 app.py', 2)
doc.add_paragraph('Main application orchestrating the workflow.')
doc.add_heading('Key Functions:', 3)
app_funcs = [
    ('main()', 'Entry point, initializes app and manages workflow'),
    ('render_stock_selection_sidebar()', 'Sidebar for category selection'),
    ('handle_file_upload()', 'Manages file upload and saved lists'),
    ('fetch_stocks_data()', 'Intelligent data fetching with mode selection')
]
for func, desc in app_funcs:
    p = doc.add_paragraph()
    p.add_run(func + ': ').bold = True
    p.add_run(desc)

doc.add_heading('Application Flow:', 3)
flow = ['1. Apply CSS', '2. Initialize session state', '3. Render header', '4. Render ticker',
        '5. Display sidebar', '6. Fetch stock list', '7. Apply sorting', '8. Fetch data',
        '9. Create DataFrame', '10. Paginate', '11. Display table', '12. Show performers']
for step in flow:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('Fetching Logic:', 3)
doc.add_paragraph('Smart mode selection based on dataset size:')
doc.add_paragraph('• >100 stocks: Bulk mode (3 workers, aggressive caching)')
doc.add_paragraph('• 50-100 stocks: Parallel mode (3 workers)')
doc.add_paragraph('• <50 stocks: Sequential mode')

doc.add_page_break()

# data_fetchers.py
doc.add_heading('4.3 data_fetchers.py', 2)
doc.add_paragraph('Handles all external API interactions.')
doc.add_heading('Key Functions:', 3)
data_funcs = [
    ('fetch_nse_csv_list()', 'Fetches from NSE CSV archives (24hr cache)'),
    ('get_index_performance()', 'Index performance (5min cache)'),
    ('get_stock_performance()', 'Core stock data fetching'),
    ('get_commodities_prices()', 'Oil, Gold, Silver, BTC, USD/INR'),
    ('fetch_stocks_bulk()', 'Bulk fetching for 100+ stocks'),
    ('get_stock_list()', 'Main function with fallback')
]
for func, desc in data_funcs:
    p = doc.add_paragraph()
    p.add_run(func + ': ').bold = True
    p.add_run(desc)

doc.add_heading('Stock Performance Logic:', 3)
doc.add_paragraph('1. Retry logic: 3 attempts with exponential backoff (3s, 9s, 27s)')
doc.add_paragraph('2. Fetch 4 months of historical data')
doc.add_paragraph('3. Get semi-live current price from info dict')
doc.add_paragraph('4. Calculate historical prices:')
doc.add_paragraph('   • 1 Week: 5 trading days back', style='List Bullet')
doc.add_paragraph('   • 1 Month: 30 calendar days back', style='List Bullet')
doc.add_paragraph('   • 2 Months: 60 days back', style='List Bullet')
doc.add_paragraph('   • 3 Months: 90 days back', style='List Bullet')
doc.add_paragraph('5. Calculate percentage changes')
doc.add_paragraph('6. Format and return result')

doc.add_heading('NSE CSV Strategy:', 3)
doc.add_paragraph('To avoid API 421 errors:')
doc.add_paragraph('1. Set browser-like headers')
doc.add_paragraph('2. Warm-up: Visit homepage to set cookies')
doc.add_paragraph('3. Fetch CSV from archives')
doc.add_paragraph('4. Parse and validate (min 40 stocks)')

doc.add_page_break()

# ui_components.py
doc.add_heading('4.4 ui_components.py', 2)
doc.add_paragraph('Streamlit UI rendering functions.')
doc.add_heading('Key Functions:', 3)
ui_funcs = [
    ('render_header()', 'Title, time, commodity prices'),
    ('render_market_indices()', 'Major and sectoral indices'),
    ('render_live_ticker()', 'Animated rolling ticker'),
    ('render_gainer_loser_banner()', 'Top gainer/loser'),
    ('render_top_bottom_performers()', 'Top 3 and bottom 3'),
    ('render_averages()', '1-year index performance'),
    ('render_pagination_controls()', 'Pagination UI')
]
for func, desc in ui_funcs:
    p = doc.add_paragraph()
    p.add_run(func + ': ').bold = True
    p.add_run(desc)

doc.add_heading('Live Ticker:', 3)
doc.add_paragraph('1. Fetch data (60s cache)')
doc.add_paragraph('2. Sort alphabetically')
doc.add_paragraph('3. Duplicate for infinite scroll')
doc.add_paragraph('4. CSS animation (120s duration)')
doc.add_paragraph('5. Hover to pause')

doc.add_page_break()

# utils.py
doc.add_heading('4.5 utils.py', 2)
doc.add_paragraph('Helper functions for formatting and processing.')
doc.add_heading('Key Functions:', 3)
util_funcs = [
    ('color_percentage()', 'Color-coded HTML for percentages'),
    ('get_current_times()', 'IST and EDT times'),
    ('format_time_display()', 'Format header display'),
    ('create_html_table()', 'HTML table with colors'),
    ('get_ticker_data()', 'Live ticker data (60s cache)')
]
for func, desc in util_funcs:
    p = doc.add_paragraph()
    p.add_run(func + ': ').bold = True
    p.add_run(desc)

doc.add_heading('Color Logic:', 3)
doc.add_paragraph('• Positive: Green (#00ff00)')
doc.add_paragraph('• Negative: Red (#ff4444)')
doc.add_paragraph('• Zero: White (#ffffff)')

doc.add_page_break()

# cache_manager.py
doc.add_heading('4.6 cache_manager.py', 2)
doc.add_paragraph('Persistent caching with Pickle.')
doc.add_heading('Features:', 3)
doc.add_paragraph('• Single-file Pickle cache (25x faster than JSON)')
doc.add_paragraph('• 6-hour expiry')
doc.add_paragraph('• Bulk operations')
doc.add_paragraph('• Automatic validation')

doc.add_heading('Key Functions:', 3)
cache_funcs = [
    ('save_to_cache()', 'Save single stock'),
    ('load_from_cache()', 'Load single stock'),
    ('save_bulk_cache()', 'Save multiple stocks'),
    ('load_bulk_cache()', 'Load multiple, return cached and missing'),
    ('clear_cache()', 'Remove cache file'),
    ('get_cache_stats()', 'Total, valid, expired counts')
]
for func, desc in cache_funcs:
    p = doc.add_paragraph()
    p.add_run(func + ': ').bold = True
    p.add_run(desc)

doc.add_page_break()

# file_manager.py
doc.add_heading('4.7 file_manager.py', 2)
doc.add_paragraph('Manages custom stock lists.')
doc.add_heading('Key Functions:', 3)
file_funcs = [
    ('ensure_saved_lists_dir()', 'Create directory'),
    ('save_list_to_csv()', 'Save list'),
    ('load_list_from_csv()', 'Load list'),
    ('delete_list_csv()', 'Delete list'),
    ('load_all_saved_lists()', 'Load all on startup')
]
for func, desc in file_funcs:
    p = doc.add_paragraph()
    p.add_run(func + ': ').bold = True
    p.add_run(desc)

doc.add_page_break()

# 5. Data Flow
doc.add_heading('5. Data Flow & Logic', 1)
doc.add_paragraph('Complete workflow:')
doc.add_paragraph('1. User opens app → main() initializes')
doc.add_paragraph('2. Load session state and saved lists')
doc.add_paragraph('3. Apply CSS styling')
doc.add_paragraph('4. Render UI (header, ticker, indices)')
doc.add_paragraph('5. User selects category')
doc.add_paragraph('6. Fetch stock list (NSE CSV or fallback)')
doc.add_paragraph('7. User clicks fetch')
doc.add_paragraph('8. Check cache, load cached, fetch missing')
doc.add_paragraph('9. Calculate performance metrics')
doc.add_paragraph('10. Create DataFrame, sort, paginate')
doc.add_paragraph('11. Display HTML table with colors')
doc.add_paragraph('12. Show top/bottom performers')

doc.add_page_break()

# 6. Key Features
doc.add_heading('6. Key Features', 1)
features = [
    'Multi-category support (Nifty 50, Next 50, Custom, Upload)',
    'Real-time data with ~15min delay',
    '1-week to 3-month performance metrics',
    'Market indices (major and sectoral)',
    'Live animated ticker (50 stocks)',
    'Commodity prices (Oil, Gold, Silver, BTC, USD/INR)',
    'Custom list upload (.txt/.csv)',
    'Persistent 6-hour cache',
    'Automatic parallel fetching',
    'Pagination (10 per page)',
    'Dark theme with responsive fonts',
    'Top/bottom performers',
    'Flexible sorting',
    'Cache management UI'
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# 7. Performance Optimizations
doc.add_heading('7. Performance Optimizations', 1)
doc.add_heading('Caching:', 2)
doc.add_paragraph('• Pickle format (25x faster than JSON)')
doc.add_paragraph('• Single-file atomic operations')
doc.add_paragraph('• 6-hour expiry')
doc.add_paragraph('• Bulk load/save')

doc.add_heading('Parallel Fetching:', 2)
doc.add_paragraph('• 3 workers for 50-100 stocks')
doc.add_paragraph('• 3 workers for 100+ stocks (bulk mode)')
doc.add_paragraph('• Exponential backoff for rate limits')
doc.add_paragraph('• Timeout handling (30s)')

doc.add_heading('Streamlit Caching:', 2)
doc.add_paragraph('• @st.cache_data for stock lists (24hr)')
doc.add_paragraph('• @st.cache_data for indices (5min)')
doc.add_paragraph('• @st.cache_data for ticker (60s)')

doc.add_heading('UI Optimizations:', 2)
doc.add_paragraph('• CSS animations (GPU-accelerated)')
doc.add_paragraph('• Pagination (10 items)')
doc.add_paragraph('• Lazy loading')
doc.add_paragraph('• Responsive font sizing')

doc.add_page_break()

# 8. Usage Guide
doc.add_heading('8. Usage Guide', 1)
doc.add_heading('Installation:', 2)
doc.add_paragraph('1. Clone repository')
doc.add_paragraph('2. Create virtual environment: python3 -m venv venv')
doc.add_paragraph('3. Activate: source venv/bin/activate')
doc.add_paragraph('4. Install: pip install -r requirements.txt')

doc.add_heading('Running:', 2)
doc.add_paragraph('streamlit run app.py')

doc.add_heading('Using the App:', 2)
doc.add_paragraph('1. Select category (Nifty 50, Next 50, Custom, Upload)')
doc.add_paragraph('2. Choose sorting (3 Months %, 2 Months %, etc.)')
doc.add_paragraph('3. Enable parallel fetching for speed')
doc.add_paragraph('4. View results with pagination')
doc.add_paragraph('5. Check top/bottom performers')
doc.add_paragraph('6. Manage cache as needed')

doc.add_heading('Uploading Custom Lists:', 2)
doc.add_paragraph('1. Prepare .txt or .csv file (one symbol per line)')
doc.add_paragraph('2. Add .NS for NSE or .BO for BSE')
doc.add_paragraph('3. Upload via sidebar')
doc.add_paragraph('4. Select exchange (Auto-detect, NSE, BSE)')
doc.add_paragraph('5. Name and save list')
doc.add_paragraph('6. Load anytime from saved lists')

doc.add_page_break()

# Dependencies
doc.add_heading('9. Dependencies', 1)
deps = ['streamlit - Web framework', 'pandas - Data manipulation', 'yfinance - Stock data',
        'requests - HTTP requests', 'pytz - Timezone handling', 'beautifulsoup4 - HTML parsing',
        'lxml - XML processing', 'plotly - Visualizations', 'curl-cffi - HTTP client']
for dep in deps:
    doc.add_paragraph(dep, style='List Bullet')

# Save document
doc.save('/Users/krishnashukla/Desktop/NSE/CascadeProjects/windsurf-project/NSE_Stock_Tracker_Documentation.docx')
print("Documentation created successfully!")
